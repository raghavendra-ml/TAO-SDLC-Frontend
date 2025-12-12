"""
Document Parser Service
Handles parsing of various document formats (Excel, Word, Text, CSV)
"""
import os
from typing import Dict, List
import openpyxl
from docx import Document
import csv

class DocumentParser:
    """Service to parse various document formats"""
    
    def __init__(self):
        self.supported_formats = ['.xlsx', '.xls', '.docx', '.doc', '.txt', '.csv']
    
    def parse_document(self, file_path: str, filename: str) -> Dict:
        """
        Parse a document and extract text content
        
        Args:
            file_path: Path to the temporary file
            filename: Original filename
            
        Returns:
            Dictionary with parsed content
        """
        ext = os.path.splitext(filename)[1].lower()
        
        if ext in ['.xlsx', '.xls']:
            return self._parse_excel(file_path)
        elif ext in ['.docx', '.doc']:
            return self._parse_word(file_path)
        elif ext == '.txt':
            return self._parse_text(file_path)
        elif ext == '.csv':
            return self._parse_csv(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _parse_excel(self, file_path: str) -> Dict:
        """Parse Excel files"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = {
                    'sheet_name': sheet_name,
                    'rows': []
                }
                
                for row in sheet.iter_rows(values_only=True):
                    # Filter out empty rows
                    row_data = [str(cell) if cell is not None else '' for cell in row]
                    if any(row_data):  # Only add non-empty rows
                        sheet_data['rows'].append(row_data)
                
                if sheet_data['rows']:
                    content.append(sheet_data)
            
            return {
                'type': 'excel',
                'filename': os.path.basename(file_path),
                'content': content,
                'text': self._extract_text_from_excel(content)
            }
        except Exception as e:
            raise Exception(f"Error parsing Excel file: {str(e)}")
    
    def _extract_text_from_excel(self, content: List[Dict]) -> str:
        """Extract plain text from Excel structure"""
        text_parts = []
        for sheet in content:
            text_parts.append(f"Sheet: {sheet['sheet_name']}")
            for row in sheet['rows']:
                text_parts.append(' | '.join(row))
        return '\n'.join(text_parts)
    
    def _parse_word(self, file_path: str) -> Dict:
        """Parse Word documents"""
        try:
            doc = Document(file_path)
            paragraphs = []
            tables = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
            
            # Combine all text
            full_text = '\n'.join(paragraphs)
            if tables:
                full_text += '\n\nTables:\n'
                for table in tables:
                    for row in table:
                        full_text += ' | '.join(row) + '\n'
            
            return {
                'type': 'word',
                'filename': os.path.basename(file_path),
                'paragraphs': paragraphs,
                'tables': tables,
                'text': full_text
            }
        except Exception as e:
            raise Exception(f"Error parsing Word document: {str(e)}")
    
    def _parse_text(self, file_path: str) -> Dict:
        """Parse plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            
            return {
                'type': 'text',
                'filename': os.path.basename(file_path),
                'lines': lines,
                'text': text
            }
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
            
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            
            return {
                'type': 'text',
                'filename': os.path.basename(file_path),
                'lines': lines,
                'text': text
            }
        except Exception as e:
            raise Exception(f"Error parsing text file: {str(e)}")
    
    def _parse_csv(self, file_path: str) -> Dict:
        """Parse CSV files"""
        try:
            rows = []
            with open(file_path, 'r', encoding='utf-8') as f:
                csv_reader = csv.reader(f)
                for row in csv_reader:
                    if any(row):  # Only add non-empty rows
                        rows.append(row)
            
            # Convert to text
            text = '\n'.join([' | '.join(row) for row in rows])
            
            return {
                'type': 'csv',
                'filename': os.path.basename(file_path),
                'rows': rows,
                'text': text
            }
        except Exception as e:
            raise Exception(f"Error parsing CSV file: {str(e)}")

